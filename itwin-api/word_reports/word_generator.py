import os
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Ordered Russian labels for ops acts (gid6-style; unknown keys appended at end).
_OPS_FIELDS: dict[str, list[tuple[str, str]]] = {
    "shurf": [
        ("nomer_akta", "Номер акта"),
        ("lineid", "ID участка"),
        ("data_utverzhdeniya_plana_shurfovok", "Дата утверждения плана шурфовок"),
        ("data_nachala_plan", "Начало (план)"),
        ("data_okonchaniya_plan", "Окончание (план)"),
        ("data_nachala", "Начало (факт)"),
        ("data_okonchaniya", "Окончание (факт)"),
        ("nomer_doma", "Номер дома"),
        ("rasstoyanie_do_blizhajshej_kamery", "Расстояние до ближайшей камеры, м"),
        ("dlina_osmotra", "Длина осмотра, м"),
        ("glubina_zalozheniya", "Глубина заложения, м"),
        ("podtoplenie_do_truby", "Подтопление до трубы"),
        ("rasstoyanie_do_relsov", "Расстояние до рельсов, м"),
        ("mesto_kontrolnoj_vyrezki_truboprovoda", "Место контрольной вырезки"),
        ("rezultaty_vyrezki", "Результаты вырезки"),
        ("rezultaty_osmotra", "Результаты осмотра"),
        ("namechennye_meropriyatiya", "Намеченные мероприятия"),
        ("meropriyatiya_po_vosstanovleniyu_prokladki", "Мероприятия по восстановлению прокладки"),
        ("fio_utverzhdaemogo", "ФИО утверждающего"),
        ("primechanie", "Примечание"),
    ],
    "osmotr": [
        ("nomer_akta", "Номер акта"),
        ("name", "Наименование"),
        ("data_osmotra", "Дата осмотра"),
        ("predpolagaemye_prichiny_razrusheniya_izolyacii_korrozii", "Предполагаемые причины"),
        ("rezultaty_osmotra", "Результаты осмотра"),
        ("namechennye_meropriyatiya", "Намеченные мероприятия"),
        ("meropriyatiya_po_vosstanovleniyu_prokladki", "Мероприятия по восстановлению"),
        ("podrazdelenie_provodivshee_raboty", "Подразделение"),
        ("fio_utverzhdaemogo", "ФИО утверждающего"),
        ("fio_1", "ФИО 1"),
        ("dolzhnost_1", "Должность 1"),
        ("fio_2", "ФИО 2"),
        ("dolzhnost_2", "Должность 2"),
        ("primechanie", "Примечание"),
    ],
    "remont": [
        ("nomer_prikaza", "Номер приказа"),
        ("data_prikaza_vvoda_v_ekspluataciyu", "Дата приказа ввода"),
        ("data_osmotra", "Дата осмотра"),
        ("vremya_osmotra", "Время осмотра"),
        ("data_utverzhdeniya_plana", "Дата утверждения плана"),
        ("data_nachala_plan", "Начало (план)"),
        ("data_okonchaniya_plan", "Окончание (план)"),
        ("data_nachala_remonta", "Начало ремонта"),
        ("data_zaversheniya_remonta", "Завершение ремонта"),
        ("otchet_po_defektu", "Отчёт по дефекту"),
        ("rezultaty_remonta", "Результаты ремонта"),
        ("len_tube_cur", "Длина трубы (факт), м"),
        ("len_channel_cur", "Длина канала (факт), м"),
        ("len_izol_cur", "Длина изоляции (факт), м"),
        ("kolichestvo_nedootpushchennoj_teplovoj_energii", "Недоотпущенная тепловая энергия"),
        ("kolichestvo_otklyuchennyh_potrebitelej", "Отключённые потребители"),
        ("primechanie", "Примечание"),
    ],
    "opres": [
        ("name", "Наименование"),
        ("istochnik_tepla", "Источник тепла"),
        ("opisaniye_kontura", "Описание контура"),
        ("date_opres", "Дата опрессовки"),
        ("vremya_provedeniya_opressovki", "Время проведения"),
        ("prodolzhitelnost_opressovki", "Продолжительность"),
        ("davlenie_opressovki_1_etap", "Давление 1 этап"),
        ("davlenie_opressovki_2_etap", "Давление 2 этап"),
        ("temperatura_raskholazhivaniya_kontura", "t расхолаживания контура"),
        ("granitsa_razdela", "Граница раздела"),
        ("defects", "Дефекты"),
        ("otchet", "Отчёт"),
        ("fio_utverzhdaemogo", "ФИО утверждающего"),
        ("fio_rukovoditel_ispytanij", "Руководитель испытаний"),
        ("primechanie", "Примечание"),
    ],
}

_SKIP_KEYS = {"shape", "geom", "geometry", "wkb_geometry", "the_geom"}


def _journal_key(journal: str) -> str:
    j = journal.lower()
    if j in {"shurf", "shurfy"}:
        return "shurf"
    if j in {"osmotr", "inspection"}:
        return "osmotr"
    if j in {"remont", "repair", "remont2"}:
        return "remont"
    if j in {"opres", "pressure-test"}:
        return "opres"
    return j


def _fmt(val: Any) -> str:
    if val is None:
        return ""
    return str(val)


def generate_defect_map_word(defect_id: int, defect_data: dict, out_dir: str = "files") -> str:
    """Генерирует .docx Карту повреждаемости (аналог KartaPovrezhdaemosti2.cpp)"""
    os.makedirs(out_dir, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("КАРТА НАРУШЕНИЯ")
    run.bold = True
    run.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run("Заявка №________\nДата заявки №________")
    run2.bold = True
    run2.font.size = Pt(12)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows_data = [
        ("Дата обнаружения:", str(defect_data.get("data_osmotra", ""))),
        ("Повреждение:", str(defect_data.get("tip_povrezhdenia", ""))),
        ("Место расположения на трубопроводе, часов:", str(defect_data.get("tsentr_povrezhdenia", ""))),
        ("Характер повреждения:", str(defect_data.get("harakter_povrezhdenia", ""))),
        ("Поврежденный трубопровод:", str(defect_data.get("povrezhdennyi_truboprovod", ""))),
        ("Состояние наружной поверхности трубопровода:", str(defect_data.get("sost_naruzhnoy", ""))),
        ("Расстояние до нарушения от ближайшей камеры, м:", str(defect_data.get("rasstoyanie", ""))),
        ("Начало времени работ:", str(defect_data.get("vremya_nachala_rabot", ""))),
        ("Окончание времени работ:", str(defect_data.get("vremya_okonchaniya_rabot", ""))),
        ("Ширина повреждения:", str(defect_data.get("shirina_povrezhdenia", ""))),
        ("Высота повреждения:", str(defect_data.get("vysota_povrezhdenia", ""))),
        ("Причины нарушения:", str(defect_data.get("prichiny", ""))),
    ]

    for label, val in rows_data:
        row = table.add_row().cells
        run_l = row[0].paragraphs[0].add_run(label)
        run_l.bold = True
        row[1].text = val

    filename = f"defect_map_{defect_id}.docx"
    filepath = os.path.join(out_dir, filename)
    doc.save(filepath)
    return filepath


def generate_ops_act_word(
    journal: str, record_id: int, data: dict, out_dir: str = "files"
) -> str:
    """Word act for shurf/osmotr/remont/opres with Russian field labels."""
    os.makedirs(out_dir, exist_ok=True)
    titles = {
        "shurf": "АКТ ШУРФОВКИ",
        "osmotr": "АКТ ОСМОТРА",
        "remont": "АКТ РЕМОНТА",
        "opres": "АКТ ОПРЕССОВКИ",
    }
    key = _journal_key(journal)
    title = titles.get(key, f"АКТ ({journal})")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Запись № {record_id}")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    used: set[str] = set()
    preferred = _OPS_FIELDS.get(key, [])
    for field, label in preferred:
        if field not in data or field in _SKIP_KEYS:
            continue
        used.add(field)
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(label).bold = True
        row[1].paragraphs[0].add_run(_fmt(data.get(field)))

    # Remaining non-empty scalar fields (skip huge/binary)
    for field, val in data.items():
        if field in used or field in _SKIP_KEYS or field == "id":
            continue
        if val is None or val == "":
            continue
        if isinstance(val, (bytes, memoryview)):
            continue
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(str(field)).bold = True
        row[1].paragraphs[0].add_run(_fmt(val))

    filename = f"{journal}_{record_id}.docx"
    path = os.path.join(out_dir, filename)
    doc.save(path)
    return path
